from langchain.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from langchain_community.utilities import WikipediaAPIWrapper
import json
import os
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
import tempfile


def generate_video_title(subject, style, video_type, api_key, creativity=0.2):
    """生成视频标题"""
    title_prompt = f"""
    请为主题"{subject}"生成一个吸引人的视频标题。
    要求：
    - 风格：{style}
    - 类型：{video_type}
    - 标题要简洁、有吸引力、符合平台特点
    - 只返回标题内容，不要其他说明
    """
    
    template = ChatPromptTemplate.from_messages([('human', title_prompt)])
    model = ChatOpenAI(
        base_url='https://twapi.openai-hk.com/v1/',
        openai_api_key=api_key,
        temperature=creativity
    )
    
    chain = template | model
    return chain.invoke({'subject': subject}).content.strip()


def generate_video_tags(subject, title, style, video_type, api_key, creativity=0.2):
    """生成视频标签"""
    tags_prompt = f"""
    为以下视频生成5-8个适合的标签：
    主题：{subject}
    标题：{title}
    风格：{style}
    类型：{video_type}
    
    要求：
    - 标签要简短、准确、容易搜索
    - 用逗号分隔
    - 只返回标签，不要其他说明
    """
    
    template = ChatPromptTemplate.from_messages([('human', tags_prompt)])
    model = ChatOpenAI(
        base_url='https://twapi.openai-hk.com/v1/',
        openai_api_key=api_key,
        temperature=creativity
    )
    
    chain = template | model
    result = chain.invoke({'subject': subject, 'title': title, 'style': style, 'video_type': video_type}).content.strip()
    
    # 解析标签
    tags = [tag.strip() for tag in result.split(',') if tag.strip()]
    return tags if tags else ["AI", "科技", "教程", "分享", "干货"]


def generate_video_description(subject, title, style, video_type, api_key, creativity=0.2):
    """生成视频简介"""
    desc_prompt = f"""
    为以下视频写一段50-100字的简介：
    主题：{subject}
    标题：{title}
    风格：{style}
    类型：{video_type}
    
    要求：
    - 50-100字左右
    - 要吸引观众点击观看
    - 体现视频的价值和亮点
    - 只返回简介内容，不要其他说明
    """
    
    template = ChatPromptTemplate.from_messages([('human', desc_prompt)])
    model = ChatOpenAI(
        base_url='https://twapi.openai-hk.com/v1/',
        openai_api_key=api_key,
        temperature=creativity
    )
    
    chain = template | model
    return chain.invoke({'subject': subject, 'title': title, 'style': style, 'video_type': video_type}).content.strip()


def generate_video_script(subject, title, video_length, style, video_type, script_structure, api_key, creativity=0.2, include_hotspot=False):
    """生成视频脚本"""
    # 脚本结构映射
    structure_prompts = {
        "开头-中间-结尾": "按照【开头、中间、结尾】三段式结构",
        "引入-冲突-高潮-结局": "按照【引入、冲突、高潮、结局】四幕剧结构",
        "问题-分析-解决": "按照【问题提出、深入分析、解决方案】逻辑结构",
        "故事-道理-启发": "按照【故事叙述、道理阐述、启发总结】结构",
        "现象-原因-对策": "按照【现象描述、原因分析、对策建议】结构",
        "吐槽-分析-解决": "按照【吐槽、分析、解决】结构"
    }
    
    # 获取热点/百科信息
    background_info = ""
    if include_hotspot:
        try:
            search = WikipediaAPIWrapper(lang='zh')
            background_info = search.run(subject)
            background_info = f"\n\n背景信息参考：{background_info[:500]}..."
        except:
            background_info = ""
    
    script_prompt = f"""
    你是一位短视频频道的博主，请用{style}风格生成一份{video_type}类型的视频脚本。
    
    视频信息：
    - 主题：{subject}
    - 标题：{title}
    - 时长：{video_length}分钟
    - 风格：{style}
    - 类型：{video_type}
    - 脚本结构：{structure_prompts.get(script_structure, '按照【开头、中间、结尾】三段式结构')}
    {background_info}
    
    要求：
    1. 开头抓住眼球，中间提供干货内容，结尾有惊喜
    2. {structure_prompts.get(script_structure, '按照【开头、中间、结尾】三段式结构')}
    3. 体现{style}风格特点
    4. 符合{video_type}类型特征
    5. 内容要适合{video_length}分钟的时长
    6. 表达方式要吸引目标受众
    7. 只返回脚本内容，不要其他说明
    """
    
    template = ChatPromptTemplate.from_messages([('human', script_prompt)])
    model = ChatOpenAI(
        base_url='https://twapi.openai-hk.com/v1/',
        openai_api_key=api_key,
        temperature=creativity
    )
    
    chain = template | model
    return chain.invoke({'subject': subject, 'title': title, 'video_length': video_length, 'style': style, 'video_type': video_type, 'script_structure': script_structure}).content.strip()


def generate_video_shots(title, script_content, style, video_type, api_key, creativity=0.2):
    """生成分镜头建议"""
    shots_prompt = f"""
    为以下视频生成分镜头建议：
    标题：{title}
    风格：{style}
    类型：{video_type}
    
    脚本内容：
    {script_content[:1000]}...
    
    要求：
    - 生成5-8个分镜头建议
    - 每个镜头包含场景描述、拍摄角度、画面要素
    - 格式：镜头1：[详细描述]、镜头2：[详细描述]...
    - 每行一个镜头，用简洁明了的文字描述
    - 不要使用Markdown格式，只返回纯文本内容
    - 不要其他说明
    """
    
    template = ChatPromptTemplate.from_messages([('human', shots_prompt)])
    model = ChatOpenAI(
        base_url='https://twapi.openai-hk.com/v1/',
        openai_api_key=api_key,
        temperature=creativity
    )
    
    chain = template | model
    result = chain.invoke({'title': title, 'style': style, 'video_type': video_type}).content.strip()
    
    # 解析分镜头 - 改进解析逻辑
    shots = []
    lines = result.split('\n')
    for line in lines:
        line = line.strip()
        # 清理可能的Markdown符号
        line = line.replace('**', '').replace('*', '').replace('#', '').replace('-', '').replace('•', '')
        line = line.strip()
        
        if line and (line.startswith('镜头') or 'Shot' in line or '场景' in line or len(line) > 10):
            shots.append(line)
    
    return shots if shots else ["镜头1：主场景全景拍摄", "镜头2：特写镜头突出重点", "镜头3：结尾总结镜头"]


def generate_video_bgm(title, script_content, style, video_type, api_key, creativity=0.2):
    """生成BGM和音效建议"""
    bgm_prompt = f"""
    为以下视频生成BGM和音效建议：
    标题：{title}
    风格：{style}
    类型：{video_type}
    
    脚本内容：
    {script_content[:1000]}...
    
    要求：
    - 提供开头BGM、中间音效、结尾BGM、转场建议
    - 音乐类型要符合视频风格和情绪
    - 每行一个建议，用简洁明了的文字描述
    - 不要使用Markdown格式，只返回纯文本内容
    - 格式示例：开头BGM：轻快节奏音乐
    - 不要其他说明
    """
    
    template = ChatPromptTemplate.from_messages([('human', bgm_prompt)])
    model = ChatOpenAI(
        base_url='https://twapi.openai-hk.com/v1/',
        openai_api_key=api_key,
        temperature=creativity
    )
    
    chain = template | model
    result = chain.invoke({'title': title, 'style': style, 'video_type': video_type}).content.strip()
    
    # 解析BGM建议 - 改进解析逻辑
    bgm_suggestions = []
    lines = result.split('\n')
    for line in lines:
        line = line.strip()
        # 清理可能的Markdown符号
        line = line.replace('**', '').replace('*', '').replace('#', '').replace('-', '').replace('•', '')
        line = line.strip()
        
        if line and (any(keyword in line for keyword in ['BGM', '音效', '背景音乐', '音乐', '转场']) or len(line) > 8):
            bgm_suggestions.append(line)
    
    return bgm_suggestions if bgm_suggestions else [
        "开头BGM：轻快的节奏音乐营造氛围",
        "中间音效：适当的转场音效",
        "结尾BGM：温和的背景音乐",
        "转场建议：简洁的过渡效果"
    ]


def generate_script(subject, video_length, creativity, api_key, 
                   style="轻松幽默", video_type="讲解类", 
                   include_shots=False, include_bgm=False, 
                   include_hotspot=False, script_structure="开头-中间-结尾",
                   include_tags=False, include_description=False):
    """
    主函数：生成完整的视频脚本
    按照优先级顺序生成各个部分
    """
    try:
        # 1. 生成视频标题
        title = generate_video_title(subject, style, video_type, api_key, creativity)
        
        result = {
            'title': title,
            'subject': subject,
            'style': style,
            'type': video_type,
            'structure': script_structure,
            'duration': video_length,
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        # 2. 根据优先级生成其他内容
        # 优先级：视频标题 > 推荐标签 > 视频简介 > 视频脚本 > 分镜头建议 > BGM和音效建议
        
        if include_tags:
            result['tags'] = generate_video_tags(subject, title, style, video_type, api_key, creativity)
        
        if include_description:
            result['description'] = generate_video_description(subject, title, style, video_type, api_key, creativity)
        
        # 3. 生成视频脚本（核心功能，始终生成）
        script_content = generate_video_script(
            subject, title, video_length, style, video_type, 
            script_structure, api_key, creativity, include_hotspot
        )
        result['script'] = script_content
        
        # 4. 生成分镜头建议
        if include_shots:
            result['shots'] = generate_video_shots(title, script_content, style, video_type, api_key, creativity)
        
        # 5. 生成BGM和音效建议
        if include_bgm:
            result['bgm_suggestions'] = generate_video_bgm(title, script_content, style, video_type, api_key, creativity)
        
        return result
        
    except Exception as e:
        raise Exception(f"视频脚本生成失败: {str(e)}")





def get_style_options():
    return ["轻松幽默", "科普教育", "情感温馨", "励志激昂", "悬疑神秘", "搞笑娱乐", "专业严谨", "文艺清新", "热血激情", "治愈温暖", "尖酸刻薄"]


def get_type_options():
    return ["讲解类", "剧情类", "Vlog类", "测评类", "教程类", "访谈类", "纪录类", "娱乐类", "开箱类", "体验类", "吐槽类"]


def get_structure_options():
    """获取可选的脚本结构"""
    return ["开头-中间-结尾", "引入-冲突-高潮-结局", "问题-分析-解决", "故事-道理-启发", "现象-原因-对策", "吐槽-分析-解决"]


def save_script_history(result):
    """保存脚本到历史记录"""
    history_dir = "script_history"
    if not os.path.exists(history_dir):
        os.makedirs(history_dir)
    
    filename = f"{history_dir}/script_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    return filename


def load_script_history():
    """加载历史脚本记录"""
    history_dir = "script_history"
    if not os.path.exists(history_dir):
        return []
    
    history_files = []
    for filename in os.listdir(history_dir):
        if filename.endswith('.json'):
            filepath = os.path.join(history_dir, filename)
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    data['filename'] = filename
                    # 添加默认收藏状态
                    if 'is_favorite' not in data:
                        data['is_favorite'] = False
                    history_files.append(data)
            except:
                continue
    
    # 按时间排序，最新的在前
    history_files.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
    return history_files


def toggle_favorite_script(filename):
    """切换脚本收藏状态"""
    history_dir = "script_history"
    filepath = os.path.join(history_dir, filename)
    
    if not os.path.exists(filepath):
        return False, "文件不存在"
    
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 切换收藏状态
        data['is_favorite'] = not data.get('is_favorite', False)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        status = "已收藏" if data['is_favorite'] else "已取消收藏"
        return True, status
    except Exception as e:
        return False, f"操作失败: {str(e)}"


def delete_script_history(filename):
    """删除历史脚本"""
    history_dir = "script_history"
    filepath = os.path.join(history_dir, filename)
    
    if not os.path.exists(filepath):
        return False, "文件不存在"
    
    try:
        os.remove(filepath)
        return True, "脚本已删除"
    except Exception as e:
        return False, f"删除失败: {str(e)}"


def get_favorite_scripts():
    """获取收藏的脚本列表"""
    all_scripts = load_script_history()
    favorite_scripts = [script for script in all_scripts if script.get('is_favorite', False)]
    return favorite_scripts


def get_script_by_filename(filename):
    """根据文件名获取脚本详情"""
    history_dir = "script_history"
    filepath = os.path.join(history_dir, filename)
    
    if not os.path.exists(filepath):
        return None
    
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
            data['filename'] = filename
            if 'is_favorite' not in data:
                data['is_favorite'] = False
            return data
    except:
        return None


def export_to_word(result, filepath):
    """导出脚本为Word文档"""
    try:
        doc = Document()
        
        # 标题
        title = doc.add_heading(result['title'], 0)
        
        # 基本信息
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ['视频风格', result['style']],
            ['视频类型', result['type']],
            ['脚本结构', result['structure']],
            ['视频时长', f"{result['duration']}分钟"]
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = value
        
        # 按优先级添加内容
        if 'tags' in result:
            doc.add_heading('推荐标签', level=1)
            doc.add_paragraph(', '.join(result['tags']))
            
        if 'description' in result:
            doc.add_heading('视频简介', level=1)
            doc.add_paragraph(result['description'])
        
        # 脚本内容
        doc.add_heading('视频脚本', level=1)
        doc.add_paragraph(result['script'])
        
        # 其他内容
        if 'shots' in result:
            doc.add_heading('分镜头建议', level=1)
            for i, shot in enumerate(result['shots'], 1):
                doc.add_paragraph(f"{i}. {shot}")
        
        if 'bgm_suggestions' in result:
            doc.add_heading('BGM音效建议', level=1)
            for bgm in result['bgm_suggestions']:
                doc.add_paragraph(f"• {bgm}")
        
        doc.save(filepath)
        return True
    except Exception as e:
        print(f"导出Word失败: {e}")
        return False


def export_to_pdf(result, filepath):
    """导出脚本为PDF文件"""
    try:
        # 处理emoji和特殊字符的函数
        def clean_text_for_pdf(text):
            """清理文本中的emoji和特殊字符，替换为文字描述"""
            if not text:
                return ""
            
            # 更全面的emoji替换映射
            emoji_replacements = {
                # 视频相关
                '🎬': '[视频]', '📝': '[文档]', '🎥': '[摄像]', '🎵': '[音乐]', 
                '🏷️': '[标签]', '📄': '[页面]', '🎞️': '[胶片]', '🎤': '[麦克风]',
                # 常用表情
                '🚀': '[火箭]', '✨': '[星星]', '💡': '[灯泡]', '🔥': '[火焰]', 
                '⭐': '[星星]', '❤️': '[爱心]', '👍': '[赞]', '👎': '[踩]', 
                '😀': '[笑脸]', '😢': '[哭脸]', '🎯': '[目标]', '📊': '[图表]', 
                '💰': '[金钱]', '🎉': '[庆祝]', '📚': '[书籍]', '🔍': '[搜索]', 
                '⚡': '[闪电]', '🌟': '[明星]', '💎': '[钻石]', '🎪': '[马戏团]', 
                '🎭': '[戏剧]', '🎨': '[艺术]', '🎧': '[耳机]', '📱': '[手机]', 
                '💻': '[电脑]', '🌈': '[彩虹]', '🌙': '[月亮]', '☀️': '[太阳]', 
                '⚠️': '[警告]', '✅': '[对勾]', '❌': '[错误]', '🔔': '[铃铛]',
                # 表情符号
                '😊': '[微笑]', '😂': '[大笑]', '🤔': '[思考]', '😍': '[心形眼]',
                '😭': '[大哭]', '🥺': '[可怜]', '😅': '[苦笑]', '🙄': '[翻白眼]',
                '😤': '[生气]', '😱': '[惊讶]', '🤗': '[拥抱]', '🤩': '[崇拜]',
                # 手势
                '👏': '[鼓掌]', '🙏': '[祈祷]', '👌': '[OK]', '✌️': '[胜利]',
                '🤝': '[握手]', '👋': '[挥手]', '🤞': '[祈愿]', '💪': '[力量]',
                # 其他常用
                '🏆': '[奖杯]', '🎁': '[礼物]', '🌸': '[樱花]', '🍀': '[四叶草]',
                '🎈': '[气球]', '🎊': '[拉花]', '🎃': '[南瓜]', '🍰': '[蛋糕]',
                '⏰': '[闹钟]', '📅': '[日历]', '📍': '[位置]', '🔗': '[链接]',
                # 网络用语相关
                '💯': '[100分]', '🔞': '[禁止]', '📢': '[广播]', '📣': '[喇叭]',
                '🏃': '[跑步]', '🛍️': '[购物]', '💸': '[花钱]', '🎲': '[骰子]',
            }
            
            # 先替换已知的emoji
            for emoji, replacement in emoji_replacements.items():
                text = text.replace(emoji, replacement)
            
            # 移除其他emoji字符（更全面的Unicode范围）
            import re
            emoji_pattern = re.compile("["
                                     u"\U0001F600-\U0001F64F"  # emoticons
                                     u"\U0001F300-\U0001F5FF"  # symbols & pictographs
                                     u"\U0001F680-\U0001F6FF"  # transport & map symbols
                                     u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
                                     u"\U00002600-\U000026FF"  # Miscellaneous Symbols
                                     u"\U00002700-\U000027BF"  # Dingbats
                                     u"\U0001F900-\U0001F9FF"  # Supplemental Symbols and Pictographs
                                     u"\U0001FA70-\U0001FAFF"  # Symbols and Pictographs Extended-A
                                     u"\U00002300-\U000023FF"  # Miscellaneous Technical
                                     u"\U0001F004-\U0001F0CF"  # Playing Cards
                                     u"\U0001F170-\U0001F251"  # Enclosed Alphanumeric Supplement
                                     "]+", flags=re.UNICODE)
            text = emoji_pattern.sub('', text)
            
            # 清理多余空格
            text = re.sub(r'\s+', ' ', text)
            return text.strip()
        
        # 注册中文字体（改进版）
        chinese_font = 'Helvetica'  # 默认字体
        try:
            import platform
            system = platform.system()
            font_registered = False
            
            if system == "Windows":
                # Windows系统字体路径（按优先级排序）
                font_paths = [
                    ("C:/Windows/Fonts/msyh.ttc", "Microsoft YaHei"),  # 微软雅黑
                    ("C:/Windows/Fonts/simsun.ttc", "SimSun"),  # 宋体
                    ("C:/Windows/Fonts/simhei.ttf", "SimHei"),  # 黑体
                    ("C:/Windows/Fonts/simkai.ttf", "SimKai"),  # 楷体
                    ("C:/Windows/Fonts/NotoSansCJK-Regular.ttc", "NotoSans"),  # Noto Sans
                ]
            elif system == "Darwin":  # macOS
                font_paths = [
                    ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
                    ("/System/Library/Fonts/Hiragino Sans GB.ttc", "HiraginoSans"),
                    ("/Library/Fonts/Arial Unicode.ttf", "ArialUnicode"),
                ]
            else:  # Linux
                font_paths = [
                    ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", "WQYMicroHei"),
                    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "DejaVuSans"),
                    ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", "Liberation"),
                ]
            
            # 尝试注册第一个可用的字体
            for font_path, font_name in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                        chinese_font = 'ChineseFont'
                        font_registered = True
                        print(f"成功注册字体: {font_name} ({font_path})")
                        break
                    except Exception as font_error:
                        print(f"注册字体失败 {font_name}: {font_error}")
                        continue
            
            if not font_registered:
                print("警告: 未找到合适的中文字体，使用默认字体")
                chinese_font = 'Helvetica'
                
        except Exception as e:
            print(f"字体注册过程出错: {e}")
            chinese_font = 'Helvetica'
        
        # 创建PDF文档
        doc = SimpleDocTemplate(filepath, pagesize=A4, 
                               topMargin=inch, bottomMargin=inch,
                               leftMargin=inch, rightMargin=inch)
        styles = getSampleStyleSheet()
        
        # 自定义样式，使用注册的中文字体
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        
        title_style = ParagraphStyle(
            'ChineseTitle',
            parent=styles['Heading1'],
            fontName=chinese_font,
            fontSize=20,
            alignment=TA_CENTER,
            spaceAfter=20,
            spaceBefore=10,
            textColor=colors.black
        )
        
        heading_style = ParagraphStyle(
            'ChineseHeading',
            parent=styles['Heading2'],
            fontName=chinese_font,
            fontSize=14,
            spaceBefore=15,
            spaceAfter=8,
            textColor=colors.darkblue
        )
        
        normal_style = ParagraphStyle(
            'ChineseNormal',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=11,
            leading=18,
            alignment=TA_LEFT,
            textColor=colors.black,
            leftIndent=10,
            rightIndent=10
        )
        
        table_style = ParagraphStyle(
            'ChineseTable',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=10,
            leading=14,
            alignment=TA_LEFT,
            textColor=colors.black
        )
        
        story = []
        
        # 1. 标题（清理emoji）
        clean_title = clean_text_for_pdf(result['title'])
        if clean_title:
            title = Paragraph(clean_title, title_style)
            story.append(title)
            story.append(Spacer(1, 15))
        
        # 2. 基本信息表格（清理emoji）
        info_data = [
            ['视频风格', clean_text_for_pdf(result.get('style', 'N/A'))],
            ['视频类型', clean_text_for_pdf(result.get('type', 'N/A'))],
            ['脚本结构', clean_text_for_pdf(result.get('structure', 'N/A'))],
            ['视频时长', f"{result.get('duration', 'N/A')}分钟"],
            ['生成时间', result.get('timestamp', 'N/A')]
        ]
        
        # 为表格数据应用样式
        formatted_info_data = []
        for row in info_data:
            formatted_row = []
            for cell in row:
                formatted_row.append(Paragraph(str(cell), table_style))
            formatted_info_data.append(formatted_row)
        
        info_table = Table(formatted_info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('BACKGROUND', (1, 0), (1, -1), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTNAME', (0, 0), (-1, -1), chinese_font),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # 3. 按优先级添加内容（清理emoji）
        if 'tags' in result and result.get('tags'):
            tags_title = Paragraph("推荐标签", heading_style)
            story.append(tags_title)
            clean_tags = [clean_text_for_pdf(tag) for tag in result['tags'] if tag]
            if clean_tags:
                tags_content = Paragraph(', '.join(clean_tags), normal_style)
                story.append(tags_content)
            story.append(Spacer(1, 12))
            
        if 'description' in result and result.get('description'):
            desc_title = Paragraph("视频简介", heading_style)
            story.append(desc_title)
            clean_description = clean_text_for_pdf(result['description'])
            if clean_description:
                desc_content = Paragraph(clean_description, normal_style)
                story.append(desc_content)
            story.append(Spacer(1, 12))
        
        # 4. 视频脚本（清理emoji）
        if result.get('script'):
            script_title = Paragraph("视频脚本", heading_style)
            story.append(script_title)
            clean_script = clean_text_for_pdf(result['script'])
            if clean_script:
                # 处理换行符
                script_paragraphs = clean_script.split('\n')
                for para in script_paragraphs:
                    if para.strip():
                        script_content = Paragraph(para.strip(), normal_style)
                        story.append(script_content)
                        story.append(Spacer(1, 6))
            story.append(Spacer(1, 12))
        
        # 5. 分镜头建议（清理emoji）
        if 'shots' in result and result.get('shots'):
            shots_title = Paragraph("分镜头建议", heading_style)
            story.append(shots_title)
            for i, shot in enumerate(result['shots'], 1):
                if shot:
                    clean_shot = clean_text_for_pdf(shot)
                    if clean_shot:
                        shot_content = Paragraph(f"{i}. {clean_shot}", normal_style)
                        story.append(shot_content)
                        story.append(Spacer(1, 6))
            story.append(Spacer(1, 12))
        
        # 6. BGM音效建议（清理emoji）
        if 'bgm_suggestions' in result and result.get('bgm_suggestions'):
            bgm_title = Paragraph("BGM音效建议", heading_style)
            story.append(bgm_title)
            for bgm in result['bgm_suggestions']:
                if bgm:
                    clean_bgm = clean_text_for_pdf(bgm)
                    if clean_bgm:
                        bgm_content = Paragraph(f"• {clean_bgm}", normal_style)
                        story.append(bgm_content)
                        story.append(Spacer(1, 6))
        
        # 生成PDF
        doc.build(story)
        print(f"PDF导出成功: {filepath}")
        return True
        
    except Exception as e:
        print(f"导出PDF失败: {e}")
        import traceback
        traceback.print_exc()
        return False


# 兼容原有接口的包装函数
def generate_script_simple(subject, video_length, creativity, api_key):
    """简化接口，保持向后兼容"""
    result = generate_script(subject, video_length, creativity, api_key)
    return result['title'], result['script']